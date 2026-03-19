#!/usr/bin/env python3
"""Build a clean manuscript DOCX from the text export and final metric files."""

from __future__ import annotations

import argparse
import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
FIGURE_MAP = {
    "Figure 1.": (ROOT / "manuscript/figures/final/figure1_kg_overlay.png", Inches(6.3)),
    "Figure 2.": (ROOT / "manuscript/figures/final/figure2_family_signal.png", Inches(6.3)),
    "Figure 3.": (ROOT / "manuscript/figures/final/figure3_recovery_profile.png", Inches(6.3)),
    "Figure 4.": (ROOT / "manuscript/figures/final/figure4_three_domain_summary.png", Inches(6.0)),
}


MAIN_HEADINGS = {
    "Abstract",
    "Introduction",
    "Methods",
    "Results",
    "Discussion",
    "Clinical Implications",
    "Limitations",
    "Future Directions",
    "Conclusions",
    "References",
    "Appendix: Representative Vignette Families",
}

SUBHEADINGS = {
    "Study Design and Reporting Standards",
    "Reference Causal Graph (KG1)",
    "Counterfactual Vignette Battery",
    "Model Selection and Access",
    "Prompting Protocol",
    "Response Parsing and Layered Validation",
    "Causal Graph Recovery",
    "Three-Dimensional Evaluation Framework",
    "Statistical Analysis",
    "Code and Data Availability",
    "Benchmark and Parser Lock",
    "Coverage of Expert Causal Structure",
    "Fidelity of Recovered Structure",
    "Edge-Level Signal Strength",
    "Discriminability: Signal versus Noise",
    "Integrated Model Profiles",
}


def load_json(path: Path) -> dict:
    import json

    with open(path) as f:
        return json.load(f)


def format_pct(value: float) -> str:
    return f"{value * 100:.1f}%"


def format_ci(metric: dict) -> str:
    return f"{format_pct(metric['estimate'])} ({format_pct(metric['ci_low'])} to {format_pct(metric['ci_high'])})"


def format_num_ci(metric: dict, digits: int = 2) -> str:
    return f"{metric['estimate']:.{digits}f} ({metric['ci_low']:.{digits}f} to {metric['ci_high']:.{digits}f})"

def build_table1_rows(domain_summary: dict) -> list[list[str]]:
    order = [
        ("DeepSeek-R1", "deepseek-r1"),
        ("Kimi K2.5", "kimi-k2.5"),
        ("Qwen3-Next-80B", "qwen3-next-80b-instruct"),
        ("Mistral-Small-24B", "mistral-small-24b"),
        ("Llama 3.1-8B", "llama-3.1-8b"),
    ]
    rows = [
        ["Dimension / Metric"] + [label for label, _ in order],
        ["COVERAGE"] + [""] * len(order),
        ["Soft recall"] + [format_ci(domain_summary[key]["soft_recall"]) for _, key in order],
        ["Soft precision"] + [format_ci(domain_summary[key]["soft_precision"]) for _, key in order],
        ["Soft FDR"] + [format_ci(domain_summary[key]["soft_fdr"]) for _, key in order],
        ["Hard recall (secondary)"] + [format_ci(domain_summary[key]["hard_recall"]) for _, key in order],
        ["FIDELITY"] + [""] * len(order),
        ["Direction accuracy (soft-detected edges)"] + [format_ci(domain_summary[key]["soft_direction_accuracy"]) for _, key in order],
        ["Direction accuracy (hard-detected edges)"] + [format_ci(domain_summary[key]["hard_direction_accuracy"]) for _, key in order],
        ["SID rate"] + [format_ci(domain_summary[key]["sid_rate"]) for _, key in order],
        ["DISCRIMINABILITY"] + [""] * len(order),
        ["SNR"] + [format_num_ci(domain_summary[key]["snr"]) for _, key in order],
        ["Detection power"] + [format_ci(domain_summary[key]["detection_power"]) for _, key in order],
    ]
    return rows


def build_table2_rows(domain_summary: dict) -> list[list[str]]:
    order = [
        ("DeepSeek-R1", "deepseek-r1"),
        ("Kimi K2.5", "kimi-k2.5"),
        ("Qwen3-Next-80B", "qwen3-next-80b-instruct"),
        ("Mistral-Small-24B", "mistral-small-24b"),
        ("Llama 3.1-8B", "llama-3.1-8b"),
    ]
    return [
        ["Measure"] + [label for label, _ in order],
        ["Mean causal JSD"] + [f"{domain_summary[key]['mean_causal_jsd']:.3f}" for _, key in order],
        ["Mean null JSD"] + [f"{domain_summary[key]['mean_null_jsd']:.3f}" for _, key in order],
        ["Null 95th percentile"] + [f"{domain_summary[key]['null_jsd_95']:.3f}" for _, key in order],
        ["Veridical split-half (aux.)"] + [format_pct(domain_summary[key]["veridical_split_half"]) for _, key in order],
        ["Risk-weighted wrong-direction rate"] + [format_pct(domain_summary[key]["weighted_wrong_direction_rate"]) for _, key in order],
        ["Risk-weighted SID rate"] + [format_pct(domain_summary[key]["weighted_sid_rate"]) for _, key in order],
    ]


def build_table3_rows() -> list[list[str]]:
    rows = [
        ["Model", "Regime", "Core profile", "Principal deployment risk"],
        ["DeepSeek-R1", "Balanced", "Best overall balance of graph breadth, direction control, and signal separation.", "Still incomplete; clinically relevant edges remain missing."],
        ["Kimi K2.5", "Broad but noisy", "Broadest graph recovery and strongest causal-versus-noise separation.", "Wrong-direction and spurious adaptation remain more frequent than in DeepSeek-R1."],
        ["Qwen3-Next-80B", "Partial and brittle", "Meaningful partial causal structure with modest discriminability.", "Limited breadth and brittle high-confidence recovery."],
        ["Mistral-Small-24B", "Partial and brittle", "Intermediate coverage with moderate signal/noise separation.", "Sparse hard-confidence structure despite reasonable soft-edge directionality."],
        ["Llama 3.1-8B", "Fragmentary", "Sparse recovery and poor causal/noise discrimination.", "False adaptation to irrelevant changes and clinically unreliable graph structure."],
    ]
    return rows


def set_base_style(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(11)
    for section in document.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)


def add_table(document: Document, rows: list[list[str]]) -> None:
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = value
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(10)
                if r_idx == 0 or (c_idx == 0 and value.isupper()):
                    for run in paragraph.runs:
                        run.bold = True
    document.add_paragraph("")


def add_paragraph(document: Document, text: str, *, italic: bool = False, bold: bool = False) -> None:
    p = document.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)


def add_figure(document: Document, caption: str) -> None:
    image_path = None
    width = Inches(6.0)
    for prefix, payload in FIGURE_MAP.items():
        if caption.startswith(prefix):
            image_path, width = payload
            break
    if image_path and image_path.exists():
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(image_path), width=width)
    add_paragraph(document, caption, italic=True)
    document.add_paragraph("")


def build_doc(txt_path: Path, output_path: Path, domain_summary: dict) -> None:
    lines = txt_path.read_text().splitlines()
    doc = Document()
    set_base_style(doc)

    title = lines[0].strip()
    subtitle = lines[1].strip()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(16)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p2.add_run(subtitle)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    doc.add_paragraph("")

    in_references = False
    for raw in lines[2:]:
        line = raw.strip()
        if not line:
            continue
        if line in MAIN_HEADINGS:
            if line in {"References", "Appendix: Representative Vignette Families"}:
                doc.add_page_break()
            doc.add_heading(line, level=1)
            in_references = line == "References"
            continue
        if line in SUBHEADINGS:
            doc.add_heading(line, level=2)
            continue
        if line.startswith("Figure "):
            add_figure(doc, line)
            continue
        if line.startswith("Table 1."):
            add_paragraph(doc, line, italic=True)
            add_table(doc, build_table1_rows(domain_summary))
            continue
        if line.startswith("Table 2."):
            add_paragraph(doc, line, italic=True)
            add_table(doc, build_table2_rows(domain_summary))
            continue
        if line.startswith("Table 3."):
            add_paragraph(doc, line, italic=True)
            add_table(doc, build_table3_rows())
            continue
        if in_references and re.match(r"^\d+\.", line):
            add_paragraph(doc, line)
            continue
        add_paragraph(doc, line)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("--txt", required=True, type=Path)
    parser.add_argument("--domain-summary", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    parser.add_argument("--backup", type=Path)
    return parser.parse_args()


if __name__ == "__main__":
    args = parse_args()
    if args.backup and args.output.exists():
        args.backup.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(args.output, args.backup)
    domain_summary = load_json(args.domain_summary)
    build_doc(args.txt, args.output, domain_summary)
